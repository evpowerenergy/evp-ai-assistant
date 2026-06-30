"""
Extract plain text from knowledge-base upload formats.
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import List, Optional

from app.config import settings
from app.utils.logger import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
MIN_EXTRACTED_CHARS = 20


@dataclass
class PageText:
    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    text: str
    pages: List[PageText] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    used_ocr: bool = False


class DocumentExtractionError(Exception):
    """Raised when text cannot be extracted from the upload."""


def _normalize_filename(filename: str) -> str:
    return (filename or "").strip().lower()


def _extension(filename: str, content_type: Optional[str]) -> str:
    name = _normalize_filename(filename)
    if name.endswith(".doc") and not name.endswith(".docx"):
        return ".doc"
    for ext in SUPPORTED_EXTENSIONS:
        if name.endswith(ext):
            return ext
    ct = (content_type or "").lower()
    if "pdf" in ct:
        return ".pdf"
    if "wordprocessingml" in ct:
        return ".docx"
    if "msword" in ct:
        return ".doc"
    if "markdown" in ct:
        return ".md"
    if "text/plain" in ct:
        return ".txt"
    return ""


def _decode_text_bytes(content_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp874", "latin-1", "cp1252"):
        try:
            return content_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentExtractionError("Could not decode text file (unsupported encoding)")


def _extract_pdf(content_bytes: bytes) -> ExtractedDocument:
    try:
        import fitz  # pymupdf
    except ImportError as e:
        raise DocumentExtractionError("PDF support not installed (pymupdf)") from e

    pages: List[PageText] = []
    parts: List[str] = []
    doc = fitz.open(stream=content_bytes, filetype="pdf")
    try:
        for i, page in enumerate(doc):
            page_text = (page.get_text() or "").strip()
            pages.append(PageText(page_number=i + 1, text=page_text))
            if page_text:
                parts.append(page_text)
    finally:
        doc.close()

    full_text = "\n\n".join(parts)
    min_expected = settings.KB_MIN_CHARS_PER_PAGE * max(len(pages), 1)

    if len(full_text.strip()) < min_expected and settings.KB_ENABLE_OCR:
        ocr_text, ocr_pages = _ocr_pdf(content_bytes)
        if ocr_text.strip():
            return ExtractedDocument(
                text=ocr_text,
                pages=ocr_pages,
                metadata={"page_count": len(ocr_pages), "extractor": "ocr"},
                used_ocr=True,
            )

    if not full_text.strip():
        raise DocumentExtractionError(
            "PDF has no extractable text. If this is a scanned document, enable OCR or use a text-based PDF."
        )

    return ExtractedDocument(
        text=full_text,
        pages=pages,
        metadata={"page_count": len(pages), "extractor": "pymupdf"},
    )


def _ocr_pdf(content_bytes: bytes) -> tuple[str, List[PageText]]:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("OCR dependencies missing; skipping OCR")
        return "", []

    pages: List[PageText] = []
    parts: List[str] = []
    doc = fitz.open(stream=content_bytes, filetype="pdf")
    try:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = (pytesseract.image_to_string(img, lang="tha+eng") or "").strip()
            pages.append(PageText(page_number=i + 1, text=text))
            if text:
                parts.append(text)
    finally:
        doc.close()

    return "\n\n".join(parts), pages


def _extract_docx(content_bytes: bytes) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as e:
        raise DocumentExtractionError("DOCX support not installed (python-docx)") from e

    doc = Document(io.BytesIO(content_bytes))
    parts: List[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts)
    if len(full_text.strip()) < MIN_EXTRACTED_CHARS:
        raise DocumentExtractionError("DOCX appears empty or has too little text")

    return ExtractedDocument(
        text=full_text,
        metadata={"extractor": "python-docx", "paragraph_count": len(doc.paragraphs)},
    )


def extract_document(
    content_bytes: bytes,
    filename: str,
    content_type: Optional[str] = None,
) -> ExtractedDocument:
    """Extract text and page structure from an uploaded file."""
    if len(content_bytes) > settings.KB_MAX_FILE_BYTES:
        raise DocumentExtractionError(
            f"File too large (max {settings.KB_MAX_FILE_BYTES // (1024 * 1024)} MB)"
        )

    ext = _extension(filename, content_type)
    if ext == ".doc":
        raise DocumentExtractionError(
            "Legacy .doc format is not supported. Please save as .docx and upload again."
        )
    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentExtractionError(f"Unsupported file type. Supported: {supported}")

    if ext in (".txt", ".md"):
        text = _decode_text_bytes(content_bytes)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if len(text) < MIN_EXTRACTED_CHARS:
            raise DocumentExtractionError("Text file is empty or too short")
        return ExtractedDocument(text=text, metadata={"extractor": "text"})

    if ext == ".pdf":
        return _extract_pdf(content_bytes)

    if ext == ".docx":
        return _extract_docx(content_bytes)

    raise DocumentExtractionError("Unsupported file type")
